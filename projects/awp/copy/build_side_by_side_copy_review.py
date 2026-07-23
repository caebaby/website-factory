import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html

from build_current_copy_review import BUILD, PAGES, extract_copy


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "copy" / "Anchor_Copy_Review_Side_by_Side.docx"

NAVY = "0F2038"
BLUE = "2D5474"
BODY = "425B72"
GOLD = "B18F4D"
LINE = "D5DDE5"
PALE_BLUE = "EEF3F6"
PALE_GOLD = "FFF9ED"
WHITE = "FFFFFF"

# compact_reference_guide with named form overrides:
# - US Letter landscape; 0.60in margins for a usable two-column form
# - 9.7167in table + 0.0833in indent = 9.80in content width
# - 9.5pt table body for long side-by-side copy, 1.15 line spacing
TABLE_WIDTH_DXA = 13972
TABLE_INDENT_DXA = 140
COLUMN_WIDTHS_DXA = [6986, 6986]
CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 140, "end": 140}


GLOBAL_BLOCKS = [
    (
        "Primary navigation",
        [
            "All pages: About · Team · Who We Serve · Resources · Process · Take the Fit Check",
        ],
        "Confirm the labels and the Fit Check as the consistent primary action across the site.",
    ),
    (
        "Brand line",
        ["\"Charting the course towards your financial legacy.\""],
        "Keep it, revise it, or provide a replacement.",
    ),
    (
        "Firm description",
        ["Houston-based fiduciary wealth management for oil & gas executives, business owners, and high-net-worth families."],
        "Confirm every descriptor and audience label is accurate and approved.",
    ),
    (
        "Legal footer placeholder",
        ["[VERIFY: RIA entity name] is a registered investment advisor. Information presented is for educational purposes only and does not intend to make an offer or solicitation for the sale or purchase of any specific securities, investments, or investment strategies. Investments involve risk and unless otherwise stated, are not guaranteed. Past performance is not a guarantee of future results."],
        "Compliance-approved replacement required before launch.",
    ),
    (
        "Other current disclosure variants",
        [
            "About / Team: [VERIFY: RIA entity name and approved regulatory disclosure before launch.] © 2026 Anchor Wealth Planning.",
            "Audience pages: Educational content only. [VERIFY full regulatory disclosure before launch.]",
            "Resources / Article: Educational information only. This material is not individualized tax, legal, or investment advice. Consult the appropriate qualified professionals before acting. © 2026 Anchor Wealth Planning.",
            "Podcast: Educational information only. Podcast content is not individualized tax, legal, or investment advice. Consult the appropriate qualified professionals before acting. © 2026 Anchor Wealth Planning.",
        ],
        "Provide one compliance-approved system and identify any page-specific variants.",
    ),
    (
        "Content-page footer description",
        [
            "Houston-based coordinated wealth planning.",
            "Coordinated planning for Houston executives, business owners, and families with complex financial lives.",
        ],
        "Approve or replace this repeated footer positioning statement.",
    ),
    (
        "Footer navigation labels",
        [
            "Homepage Explore: About · Team · Services · Resources · Contact",
            "Homepage Specialties: Oil & Gas Executives · Business Owners · HNW Families · Equity Compensation · Estate & Legacy · Coordinated Planning",
            "Content pages: Explore · Start · Who We Serve · Take the Fit Check · Oil and Gas Executives",
        ],
        "Confirm the footer labels and destinations.",
    ),
    (
        "Repeated Fit Check footer",
        [
            "The Anchor Fit Check",
            "Four questions. A better first conversation.",
            "See whether your situation, timing, and coordination needs match the way Anchor works.",
            "01 Your situation · 02 Current complexity · 03 Coordination gap · 04 Best next step.",
            "CTA: Start the Fit Check.",
        ],
        "Approve the promise, steps, and CTA wording.",
    ),
]

HERO_BLOCKS = [
    (
        "Hero positioning label",
        ["Houston, TX · Fiduciary Wealth Management"],
        "Confirm the location and whether fiduciary language is approved.",
    ),
    (
        "Rotating hero — Oil & Gas",
        [
            "When your equity vests, who sees the whole picture?",
            "RSUs, deferred comp, blackout periods, commodity cycles. Anchor coordinates your investments, taxes, and estate plan with your CPA and attorney — so every vest, bonus, and decision happens in context.",
        ],
        "Approve the headline and description for the lead hero.",
    ),
    (
        "Rotating hero — Business Owners",
        [
            "Your business is your largest asset. What’s the plan for it?",
            "Entity structure, exit timing, liquidity events, and the line between business value and personal wealth. Anchor starts the planning conversation years before the LOI arrives.",
        ],
        "Approve or replace the business-owner hero copy.",
    ),
    (
        "Rotating hero — Families",
        [
            "Your wealth grew. Did your plan keep up?",
            "Trust funding, beneficiary reviews, multi-generational coordination. Anchor keeps your CPA, attorney, and advisor reading from the same page — so your estate matches your actual net worth.",
        ],
        "Approve or replace the family hero copy.",
    ),
    (
        "Hero actions and proof",
        [
            "Market tabs: 01 Oil & Gas · 02 Business Owners · 03 Families",
            "CTA: Schedule a 30-Minute Conversation",
            "CTA: Take the 4-Question Fit Check",
            "[Credential 01] · [Credential 02] · [XX]+ yrs · Houston-based",
        ],
        "Confirm the conversion path and supply verified proof.",
    ),
]


# Alex should complete these first. They are ordered by launch risk, factual
# dependency, and conversion impact—not by where they happen to appear on-site.
PRIORITY_REVIEW = [
    (
        "REQUIRED BEFORE LAUNCH",
        "Firm identity and legal disclosure",
        [
            "CURRENT POSITIONING: Houston-based fiduciary wealth management for oil & gas executives, business owners, and high-net-worth families.",
            "CURRENT PLACEHOLDER: [VERIFY: RIA entity name] is a registered investment advisor.",
            "CURRENT DISCLOSURE: Information presented is for educational purposes only and does not intend to make an offer or solicitation for the sale or purchase of any specific securities, investments, or investment strategies. Investments involve risk and unless otherwise stated, are not guaranteed. Past performance is not a guarantee of future results.",
        ],
        "Provide the exact legal entity name, approved regulatory description, and compliance-approved disclosure. Confirm whether ‘fiduciary wealth management’ is approved public language.",
    ),
    (
        "REQUIRED BEFORE LAUNCH",
        "Alex’s public title, biography, and founding story",
        [
            "CURRENT TITLE PLACEHOLDER: Alex Miller",
            "CURRENT STORY POSITIONING: Anchor was founded to coordinate the decisions that fall between investments, taxes, estate work, and business planning.",
            "CURRENT PLACEHOLDER: Alex’s approved biography, credentials, prior-firm history, community involvement, and personal details will be inserted after client and compliance review.",
        ],
        "Supply Alex’s exact title, approved short biography, prior-firm history, founding story, community involvement, and any personal details that should appear publicly.",
    ),
    (
        "REQUIRED BEFORE LAUNCH",
        "Verified experience, credentials, and accolades",
        [
            "[XX]+ years in the industry — Experience across institutional wealth management and independent advice.",
            "[Credential 01] — Approved professional designation and issuing organization go here.",
            "[Credential 02] — Approved license, certification, or specialist training goes here.",
            "[Accolade 01] — Approved recognition, year, source, and selection criteria go here.",
        ],
        "Replace every placeholder with a verified fact. Include issuing organizations, dates where relevant, and accolade source/selection criteria—or remove the item.",
    ),
    (
        "REQUIRED BEFORE LAUNCH",
        "Team structure and public profiles",
        [
            "CURRENT TEAM PROMISE: People who own the next step. A coordinated plan works only when responsibility is clear.",
            "CURRENT PLACEHOLDERS: Two additional [Name & Role] profiles for an advisor, client service lead, operations partner, or other approved team member.",
        ],
        "Identify who should appear publicly. Provide each person’s role, responsibilities, credentials, biography, headshot, and contact routing. Confirm whether unfilled profiles should remain hidden.",
    ),
    (
        "REQUIRED BEFORE LAUNCH",
        "Services Anchor actually delivers",
        [
            "Coordinated Planning — Joint planning meetings with the CPA and estate attorney; one living document shared across the team.",
            "Investment Management — Strategy built around vesting schedules, trading restrictions, commodity exposure, and family cash-flow needs.",
            "Tax Strategy & Coordination — Year-end coordination with the CPA; modeling of vesting, deferred compensation, charitable timing, and estimated payments.",
            "Equity Compensation — RSUs, NQDC elections, 10b5-1 plans, blackout periods, concentration, and tax-aware diversification.",
            "Estate & Legacy — Existing-trust review, beneficiary/funding coordination, and work with the estate attorney.",
            "Business Owner Planning — Exit readiness, entity structure, personal/enterprise balance sheets, and post-liquidity planning.",
        ],
        "Confirm each service is genuinely offered and accurately described. Correct anything that overstates Anchor’s role versus the CPA or attorney. Confirm whether joint planning meetings and one living document are standard client experiences.",
    ),
    (
        "REQUIRED BEFORE LAUNCH",
        "Fit Check questions, thresholds, and response promise",
        [
            "PROMISE: Four questions. Two minutes. We’ll tell you honestly whether we can help—and if we can’t, we’ll refer you to someone who can.",
            "SITUATION OPTIONS: Oil and gas executive with equity compensation · Business owner considering a transition · High-net-worth family with estate and legacy complexity · Other / not sure.",
            "INVESTABLE-ASSET OPTIONS: $500K–$2M · $2M–$5M · $10M+.",
            "FIRST-NEED OPTIONS: Tax coordination · Equity compensation · Business transition · Estate and legacy · Investment strategy · Not sure.",
            "CURRENT REVIEW RESPONSE: Production will route answers to a fit-results experience. We’ll review your responses and follow up within one business day. No pitch. Just a clear next step.",
        ],
        "Approve every option and promise. Resolve the missing $5M–$10M band, confirm response timing, and confirm whether Anchor will make referrals when the visitor is not a fit.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Homepage hero — Oil & Gas",
        [
            "When your equity vests, who sees the whole picture?",
            "RSUs, deferred comp, blackout periods, commodity cycles. Anchor coordinates your investments, taxes, and estate plan with your CPA and attorney—so every vest, bonus, and decision happens in context.",
        ],
        "Approve, edit, or replace this headline and description. Confirm this should be the lead hero audience.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Homepage hero — Business Owners",
        [
            "Your business is your largest asset. What’s the plan for it?",
            "Entity structure, exit timing, liquidity events, and the line between business value and personal wealth. Anchor starts the planning conversation years before the LOI arrives.",
        ],
        "Approve or replace this headline and description. Confirm the ‘before the LOI’ framing matches Alex’s process.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Homepage hero — High-Net-Worth Families",
        [
            "Your wealth grew. Did your plan keep up?",
            "Trust funding, beneficiary reviews, multi-generational coordination. Anchor keeps your CPA, attorney, and advisor reading from the same page—so your estate matches your actual net worth.",
        ],
        "Approve or replace this headline and description. Confirm the preferred public audience label.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Why Anchor",
        [
            "One advisor who sees the whole picture.",
            "Complex wealth does not need another specialist working alone. It needs one accountable lead who keeps every professional, deadline, and decision moving from the same plan.",
            "Coordinated, not siloed · One accountable lead · Built for complexity · Houston fluency.",
        ],
        "Approve the core differentiation. Clarify exactly what Alex owns, what outside professionals own, and whether ‘one accountable lead’ is the right promise.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Who Anchor is—and is not—for",
        [
            "Oil & Gas Executives — RSUs, NQDC elections, blackout periods, and commodity-cycle exposure.",
            "Business Owners — Exit readiness, entity structure, and the line between business value and personal wealth.",
            "High-Net-Worth Families — Trust funding, current beneficiaries, and a plan the next generation can understand.",
            "CURRENT FRAMING: If you don’t see yourself below, we may not be the right fit—and we’ll tell you that honestly.",
        ],
        "Define the strongest-fit client for each audience: role or situation, complexity, approximate assets, timing, and any industries or situations Alex does not serve.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Primary conversion path",
        [
            "CURRENT PRIMARY CTA: Take the Fit Check.",
            "CURRENT SECONDARY CTA: Schedule a 30-Minute Conversation.",
            "CURRENT BEHAVIOR: Call CTAs route into the Fit Check before scheduling.",
        ],
        "Choose the intended sequence: Fit Check first, direct scheduling first, or Fit Check followed by scheduling. Confirm the preferred meeting length and CTA wording.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Visitor pain language and first-person statements",
        [
            "‘No one sees the full picture. My CPA handles taxes, my advisor handles investments, my attorney has the estate docs—and none of them talk to each other.’",
            "‘My RSUs vest, I get a huge tax hit, and by the time I think about selling to diversify, I’m emotionally attached to the stock and it’s already dropped.’",
            "‘We have a trust that was set up years ago. I’m honestly not sure if it’s even funded correctly, or if it still matches what we actually own.’",
            "‘I know my business inside and out, but I have no idea what it’s actually worth or what I’d walk away with after taxes if I sold.’",
        ],
        "Approve, soften, or replace these audience statements. They should sound recognizably true to Alex’s clients without pretending to be direct client quotations unless they are sourced and approved.",
    ),
    (
        "HIGH-IMPACT MESSAGE",
        "Client process and ongoing service rhythm",
        [
            "01 Discovery — A 30-minute conversation about the situation and concerns.",
            "02 Assessment — Review investments, tax returns, equity documents, estate plan, and business interests.",
            "03 The Plan — Build the coordinated investment, tax, estate, and business-owner roadmap.",
            "04 Partnership — Implement the plan and meet quarterly to keep it current.",
        ],
        "Confirm the real client journey, documents reviewed, deliverables, meeting cadence, and what changes before or after engagement.",
    ),
    (
        "CONTENT DIRECTION",
        "Resources and Anchor Wealth Podcast",
        [
            "SERIES NAME: Anchor Wealth Podcast.",
            "DIRECTORY PROMISE: Clear thinking for decisions that touch everything.",
            "CURRENT TOPIC AREAS: Equity and tax timing · Business transition · Family and legacy · Advisor coordination.",
            "CURRENT CONTENT MIX: Articles · Podcast episodes · Checklists · Frameworks.",
            "CURRENT AUTHORITY: Alex Miller as article author and Anchor Wealth Podcast host; source-backed articles, episode notes, and transcripts lead to the Fit Check.",
        ],
        "Approve the podcast name, Alex as author/host, and topic ownership. Choose the first three to six real article or episode topics, publishing platforms, transcript versus show-note policy, compliance review process, and any guide, checklist, or newsletter offer.",
    ),
]


# Remaining copy worth a human look after the priorities above are settled.
# Routine labels, repeated CTAs, metadata, alt text, and duplicated footer copy
# are intentionally excluded from the client workload.
SECONDARY_PAGE_GROUPS = {
    "home-v6.html": [
        ("Homepage supporting message", [
            "The firm was designed around one missing role.",
            "Turn separate expertise into one coordinated course.",
            "You've been the project manager long enough.",
        ]),
    ],
    "about.html": [
        ("About-page supporting message", [
            "One plan. One accountable lead.",
            "Expert advice can still be fragmented.",
            "Keep every decision on the same course.",
        ]),
    ],
    "oil-gas-executives.html": [
        ("Oil & Gas page supporting message", [
            "Your compensation moves with the cycle.",
            "Your career and wealth may share one risk.",
            "Map the calendar",
            "Model the exposure",
            "Coordinate the professionals",
        ]),
    ],
    "business-owners.html": [
        ("Business-owner page supporting message", [
            "The exit conversation starts before the offer.",
            "The company is valuable. Is the wealth portable?",
            "Separate the balance sheets",
            "Build exit readiness",
            "Plan the other side",
        ]),
    ],
    "high-net-worth-families.html": [
        ("Family page supporting message", [
            "Keep the plan current across generations.",
            "Your wealth changed. Did the plan?",
            "Review what exists",
            "Close implementation gaps",
            "Prepare the family",
        ]),
    ],
    "resources.html": [
        ("Resources-directory introduction", [
            "Clear thinking for decisions that touch everything.",
            "Choose how you want to think it through.",
            "Start with the decision, not the format.",
        ]),
        ("Sample resource titles and descriptions", [
            "How RSU vesting creates a tax-timing trap",
            "What coordinated planning looks like before year-end",
            "What one living plan actually coordinates",
            "The exit conversation starts before the LOI",
            "Does the estate plan still match what the family owns?",
        ]),
    ],
    "blog-template.html": [
        ("Sample article — core message", [
            "How RSU vesting creates a tax-timing trap",
            "The short answer",
            "A better sequence before the next vest",
        ]),
        ("Sample article — supporting sections", [
            "What should you gather before the vest date?",
            "Common RSU tax-planning questions",
            "Primary sources and further reading",
        ]),
    ],
    "podcast-template.html": [
        ("Sample podcast episode — opening message", [
            "What coordinated planning looks like before year-end",
            "The calendar is where coordination becomes visible.",
            "Start with the decisions, not the meetings",
        ]),
        ("Sample podcast episode — practical takeaways", [
            "Give every decision one owner",
            "Move information before the deadline",
            "Update the plan after the decision",
            "Four items for the first calendar pass",
        ]),
    ],
}


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in CELL_MARGINS_DXA.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in COLUMN_WIDTHS_DXA:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, COLUMN_WIDTHS_DXA[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), LINE)


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.60)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 18, 12, 7, NAVY),
        ("Heading 2", 13, 10, 5, BLUE),
        ("Heading 3", 11, 7, 4, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    r = header.add_run("ANCHOR WEALTH PLANNING  |  COPY REVIEW WORKBOOK")
    set_font(r, size=8, color=BODY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Alex copy review  |  ")
    set_font(r, size=8, color=BODY)
    add_page_field(footer)


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_small_label(cell, text, color=GOLD):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_font(r, size=7.5, color=color, bold=True)


def add_header_row(table):
    row = table.rows[0]
    repeat_header_row(row)
    labels = ["CURRENT WEBSITE COPY", "ALEX'S VERSION / CHANGES"]
    for cell, label in zip(row.cells, labels):
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_font(r, size=9, color=WHITE, bold=True)
        set_shading(cell, NAVY)


def change_prompt(title, text, default_prompt):
    if default_prompt:
        return default_prompt
    combined = f"{title} {' '.join(text)}".lower()
    if any(token in combined for token in ["[verify", "credential", "accolade", "years", "registered investment advisor"]):
        return "Verified fact or compliance-approved replacement required."
    if any(token in combined for token in ["fit check", "cta:", "schedule a", "book a call"]):
        return "Keep the action or type the preferred CTA and next step."
    return "Keep, edit, or replace this copy."


def add_copy_row(table, title, text, prompt=None, priority=None):
    row = table.add_row()
    prevent_row_split(row)
    left, right = row.cells
    set_shading(left, PALE_GOLD if priority == "REQUIRED BEFORE LAUNCH" else WHITE)
    set_shading(right, PALE_GOLD)

    clear_cell(left)
    add_small_label(left, priority or "Secondary copy")
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, size=10.5, color=NAVY, bold=True)
    for item in text:
        if not item:
            continue
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=9.5, color=BODY)

    clear_cell(right)
    add_small_label(right, "Decision", color=BLUE)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Keep  /  Edit  /  Replace")
    set_font(r, size=9.5, color=NAVY, bold=True)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(change_prompt(title, text, prompt))
    set_font(r, size=8.5, color=BODY, italic=True)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Alex's replacement copy or notes:")
    set_font(r, size=8.5, color=BLUE, bold=True)
    response_lines = 3 if priority == "REQUIRED BEFORE LAUNCH" else 2
    for _ in range(response_lines):
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(" ")
        set_font(r, size=9.5, color=BODY)


def add_decision_row(table, prompts):
    row = table.add_row()
    prevent_row_split(row)
    left, right = row.cells
    set_shading(left, PALE_BLUE)
    set_shading(right, PALE_GOLD)
    clear_cell(left)
    add_small_label(left, "Page-level decisions", color=BLUE)
    for number, prompt in enumerate(prompts, 1):
        p = left.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{number}. {prompt}")
        set_font(r, size=8.7, color=BODY)

    clear_cell(right)
    add_small_label(right, "Alex's answers", color=BLUE)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Type answers here. Short, plain-language answers are fine.")
    set_font(r, size=8.5, color=BODY, italic=True)
    for _ in range(max(4, min(8, len(prompts)))):
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.add_run(" ")


def metadata_block_for_page(filename):
    tree = html.fromstring((BUILD / filename).read_text(encoding="utf-8"))
    values = []

    title = tree.xpath("string(//title)").strip()
    if title:
        values.append(f"SEO TITLE: {title}")

    for label, xpath in [
        ("META DESCRIPTION", "//meta[@name='description']/@content"),
        ("AUTHOR", "//meta[@name='author']/@content"),
        ("SOCIAL TITLE", "//meta[@property='og:title']/@content"),
        ("SOCIAL DESCRIPTION", "//meta[@property='og:description']/@content"),
    ]:
        result = tree.xpath(xpath)
        if result:
            value = " ".join(result[0].split())
            if value and value not in " ".join(values):
                values.append(f"{label}: {value}")

    schema_summaries = []
    for node in tree.xpath("//script[@type='application/ld+json']/text()"):
        try:
            data = json.loads(node)
        except (TypeError, ValueError):
            continue
        records = data.get("@graph", []) if isinstance(data, dict) and "@graph" in data else [data]
        for record in records:
            if not isinstance(record, dict):
                continue
            schema_type = record.get("@type")
            if isinstance(schema_type, list):
                schema_type = ", ".join(schema_type)
            headline = record.get("headline") or record.get("name")
            description = record.get("description")
            author = record.get("author")
            author_name = author.get("name") if isinstance(author, dict) else None
            date_published = record.get("datePublished")
            details = [str(value) for value in [schema_type, headline, description, author_name, date_published] if value]
            if details:
                schema_summaries.append(" · ".join(details))
    if schema_summaries:
        values.append("STRUCTURED DATA: " + " | ".join(schema_summaries))

    return (
        "SEO, social, and GEO metadata",
        values or ["No page-specific SEO metadata is currently present."],
        "Approve the search/social wording and verify every person, organization, date, and schema fact.",
    )


def accessibility_block_for_page(filename):
    tree = html.fromstring((BUILD / filename).read_text(encoding="utf-8"))
    values = []

    def add(label, value):
        value = " ".join((value or "").split())
        item = f"{label}: {value}" if value else ""
        if item and item not in values:
            values.append(item)

    for value in tree.xpath("//img/@alt"):
        add("IMAGE ALT", value)
    for value in tree.xpath("//input[@placeholder]/@placeholder"):
        add("FORM PLACEHOLDER", value)
    for element in tree.xpath("//*[contains(concat(' ', normalize-space(@class), ' '), ' sr-only ')]"):
        add("SCREEN-READER LABEL", element.text_content())
    for value in tree.xpath("//*[@aria-label]/@aria-label"):
        add("INTERFACE LABEL", value)

    if not values:
        return None
    return (
        "Accessibility and interface text",
        values,
        "Confirm image descriptions and page-specific labels; standard navigation/control labels can remain unless inaccurate.",
    )


def blocks_for_page(filename):
    blocks = []
    title = "Page introduction"
    parts = []
    pending_label = None
    pending_details = []

    def flush(force=False):
        nonlocal title, parts
        cleaned = [part for part in parts if part]
        if cleaned or (force and title not in {"Page introduction", "Copy block"}):
            blocks.append((title, cleaned))
        title = "Copy block"
        parts = []

    for tag, text, classes in extract_copy(BUILD / filename):
        if "eyebrow" in classes:
            flush()
            pending_label = text
            continue
        is_lead_detail = (
            (tag == "small" and (text.startswith("Step ") or " · " in text or text.endswith("preview")))
            or bool(classes.intersection({"chapter-number", "reel-tag", "serve-card-tag"}))
        )
        if is_lead_detail:
            flush()
            if "chapter-number" in classes:
                pending_details.append(f"CHAPTER: {text}")
            elif classes.intersection({"reel-tag", "serve-card-tag"}):
                pending_details.append(f"LABEL: {text}")
            else:
                pending_details.append(f"DETAIL LABEL: {text}")
            continue
        if tag in {"h1", "h2", "h3", "h4", "legend", "strong"}:
            flush(force=True)
            title = text
            if pending_label:
                parts.append(f"SECTION LABEL: {pending_label}")
                pending_label = None
            if pending_details:
                parts.extend(pending_details)
                pending_details = []
            continue

        if pending_label:
            flush()
            title = pending_label
            pending_label = None

        if tag == "label":
            parts.append(f"OPTION: {text}")
        elif tag == "li":
            parts.append(text)
        elif tag in {"a", "button"}:
            parts.append(f"CTA / LINK: {text}")
        elif tag == "small":
            parts.append(f"DETAIL LABEL: {text}")
        elif tag == "figcaption":
            parts.append(f"IMAGE CAPTION: {text}")
        elif tag == "summary":
            parts.append(f"EXPANDER LABEL: {text}")
        elif "svc-card-pain" in classes:
            parts.append(f"SUPPORTING QUOTE: {text}")
        elif "when" in classes:
            parts.append(f"TIMING LABEL: {text}")
        elif classes.intersection({"trust-status", "tag", "chip", "reel-tag", "serve-card-tag"}):
            parts.append(f"LABEL: {text}")
        else:
            parts.append(text)
    if pending_label:
        flush()
        title = pending_label
        pending_label = None
    if pending_details:
        flush()
        title = "Detail labels"
        parts.extend(pending_details)
        pending_details = []
    flush(force=True)
    return blocks


def secondary_groups_for_page(filename):
    lookup = {}
    for title, text in blocks_for_page(filename):
        lookup.setdefault(title, text)

    groups = []
    for group_title, block_titles in SECONDARY_PAGE_GROUPS.get(filename, []):
        combined = []
        for block_title in block_titles:
            text = lookup.get(block_title)
            if text is None:
                raise ValueError(f"Missing secondary copy block in {filename}: {block_title}")
            combined.append(block_title)
            for item in text:
                if item.startswith(("SECTION LABEL:", "CTA / LINK:", "LABEL:", "DETAIL LABEL:", "CHAPTER:")):
                    continue
                if item.startswith("Home /") or item in {
                    "Articles", "Podcasts", "Frameworks", "Listen to the podcast",
                    "View framework", "Explore topic",
                }:
                    continue
                combined.append(item)
        groups.append((group_title, combined))
    return groups


def add_page_heading(doc, kicker, title, url=None, page_break_before=False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(kicker.upper())
    set_font(r, size=8, color=GOLD, bold=True)
    doc.add_paragraph(title, style="Heading 1")
    if url:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(url)
        set_font(r, size=8, color=BLUE)


def new_review_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table)
    set_table_borders(table)
    add_header_row(table)
    return table


def build():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("CLIENT DECISION WORKBOOK")
    set_font(r, size=9, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Anchor Website Copy Priorities")
    set_font(r, size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Part 1 contains the decisions that materially affect accuracy, trust, positioning, and conversion. Complete those first. Part 2 contains secondary wording that only needs attention if Alex wants it changed.")
    set_font(r, size=11, color=BODY)

    instruction_table = new_review_table(doc)
    add_copy_row(
        instruction_table,
        "How to complete this review",
        [
            "1. Complete every REQUIRED BEFORE LAUNCH block.",
            "2. Review the HIGH-IMPACT MESSAGE blocks and change only what does not sound like Alex.",
            "3. Complete CONTENT DIRECTION if articles or the podcast will launch with the site.",
            "4. Use Part 2 only for wording Alex actively wants to change.",
            "5. Plain-language notes are enough; final copy will be polished after the factual decisions are settled.",
            "REMOVED FROM THIS WORKBOOK: routine navigation, repeated buttons, standard footer labels, SEO/social metadata, schema, accessibility labels, and duplicated copy that can safely remain as written.",
        ],
        "Use this box for any overall voice direction or non-negotiable language.",
        priority="START HERE",
    )

    doc.add_page_break()
    add_page_heading(doc, "Part 1 — complete first", "The copy decisions that matter")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("These blocks are intentionally ordered by launch risk and business impact. Alex does not need to work through the website page by page.")
    set_font(r, size=9.5, color=BODY, italic=True)
    table = new_review_table(doc)
    for priority, title, text, prompt in PRIORITY_REVIEW:
        add_copy_row(table, title, text, prompt, priority=priority)

    add_page_heading(doc, "Part 2 — only if needed", "Secondary copy worth reviewing", page_break_before=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("The strongest strategic and factual decisions are already in Part 1. Everything below can stay as written unless Alex wants a different tone, emphasis, or example.")
    set_font(r, size=10, color=BODY)

    first_secondary_page = True
    for page_title, filename, url, _prompts in PAGES:
        selected = secondary_groups_for_page(filename)
        if not selected:
            continue
        add_page_heading(doc, "Secondary page review", page_title, url, page_break_before=not first_secondary_page)
        first_secondary_page = False
        table = new_review_table(doc)
        for title, text in selected:
            add_copy_row(table, title, text)

    doc.core_properties.title = "Anchor Website Copy Priorities — Side by Side"
    doc.core_properties.subject = "Prioritized editable client copy decision workbook"
    doc.core_properties.author = "Website Factory"
    doc.core_properties.keywords = "Anchor Wealth Planning, copy review, Alex Miller"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
