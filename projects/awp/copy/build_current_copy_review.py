from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html


ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "build"
OUTPUT = ROOT / "copy" / "Anchor_Wealth_Planning_Current_Copy_Review.docx"
AI_OUTPUT = ROOT / "copy" / "ANCHOR_AI_COPY_PACKET.md"

NAVY = "0F2038"
BLUE = "2D5474"
BODY = "425B72"
GOLD = "B18F4D"
LIGHT = "EEF2F5"
LINE = "D8DEE5"
WHITE = "FFFFFF"

PAGES = [
    (
        "Homepage",
        "home-v6.html",
        "https://caebaby.github.io/anchor-review/",
        [
            "Choose the lead hero audience and approve all three rotating hero headlines.",
            "Confirm whether “fiduciary wealth management” is accurate and compliance-approved.",
            "Approve, soften, or replace the first-person pain statements used in the ticker and problem section.",
            "Supply the approved biography, prior-firm history, years in the industry, credentials, and any accolade with source and selection criteria.",
            "Confirm every service described is offered and whether “joint planning meetings” is the normal client experience.",
            "Confirm the Fit Check asset bands. The current build moves from $2M–$5M directly to $10M+ and may need a $5M–$10M option.",
            "Decide whether the primary action is a Fit Check, a 30-minute call, or a Fit Check followed by scheduling. Current call CTAs route into the Fit Check.",
            "Approve the response promise, referral language, and “no pitch” wording.",
            "Replace preview story and short-form videos with approved Alex footage when available.",
        ],
    ),
    (
        "About",
        "about.html",
        "https://caebaby.github.io/anchor-review/about.html",
        [
            "Provide Alex’s approved founding story: why Anchor exists and what was missing in his prior environment.",
            "Confirm the promise that Anchor maintains one living plan and coordinates outside professionals.",
            "Approve the phrase “one accountable lead” and clarify what Alex owns versus what the CPA or attorney owns.",
            "Supply approved credentials, experience, community involvement, and personal details that should appear publicly.",
        ],
    ),
    (
        "Team",
        "team.html",
        "https://caebaby.github.io/anchor-review/team.html",
        [
            "Confirm Alex’s exact public title.",
            "Provide Alex’s approved short biography and headshot.",
            "Identify every team member or outside partner who should appear, including role, responsibilities, credentials, biography, headshot, and contact routing.",
            "Decide whether placeholders should remain hidden until real team profiles are approved.",
        ],
    ),
    (
        "Oil & Gas Executives",
        "oil-gas-executives.html",
        "https://caebaby.github.io/anchor-review/oil-gas-executives.html",
        [
            "Confirm the audience language accurately reflects Alex’s ideal Houston energy client.",
            "Verify that Anchor advises on RSUs, deferred compensation, NQDC elections, trading windows, and concentrated employer exposure.",
            "Provide one anonymized real-world planning situation that can replace generic explanation with specific proof.",
            "Flag any compensation or tax language that requires compliance qualification.",
        ],
    ),
    (
        "Business Owners",
        "business-owners.html",
        "https://caebaby.github.io/anchor-review/business-owners.html",
        [
            "Define the ideal owner profile by company size, industry, transition horizon, and investable assets.",
            "Confirm whether valuation, entity structure, succession, pre-sale planning, and post-liquidity planning are all within scope.",
            "Approve or replace the “before the LOI” framing.",
            "Provide one anonymized owner situation that shows Alex’s coordination role without making an unverified performance claim.",
        ],
    ),
    (
        "High-Net-Worth Families",
        "high-net-worth-families.html",
        "https://caebaby.github.io/anchor-review/high-net-worth-families.html",
        [
            "Confirm whether “high-net-worth families” is the preferred public audience label.",
            "Verify the scope around trust funding, beneficiary reviews, account ownership, family communication, and estate-attorney coordination.",
            "Clarify the typical family complexity or asset level that signals a strong fit.",
            "Provide one anonymized family situation to make the page more specific and credible.",
        ],
    ),
    (
        "Resources Directory",
        "resources.html",
        "https://caebaby.github.io/anchor-review/resources.html",
        [
            "Approve “Anchor Wealth Podcast” as the public podcast-series name.",
            "Choose the initial three to six article and podcast topics Alex can credibly own.",
            "Confirm whether podcast links should point to Spotify, Apple Podcasts, YouTube, or an embedded player once episodes are live.",
            "Approve the topic taxonomy: equity and tax timing, business transition, family and legacy, and advisor coordination.",
            "Identify any downloadable guide, checklist, or newsletter signup that belongs in the directory.",
        ],
    ),
    (
        "Article Template",
        "blog-template.html",
        "https://caebaby.github.io/anchor-review/blog-template.html",
        [
            "Confirm Alex as author and approve the displayed publication date and reading time.",
            "Have compliance review the RSU, supplemental-wage withholding, trading-window, tax, and planning statements.",
            "Confirm the article matches Alex’s advice process and terminology.",
            "Approve the cited IRS sources and add any employer-plan, legal, or compliance caveat Alex normally uses.",
            "Choose a production article URL, author biography, social-share image, and final call to action before launch.",
        ],
    ),
    (
        "Podcast Episode Template",
        "podcast-template.html",
        "https://caebaby.github.io/anchor-review/podcast-template.html",
        [
            "Confirm Alex as host and approve “Anchor Wealth Podcast” as the series title.",
            "Replace preview language with the final episode number, date, duration, audio/video embed, platform links, and guest information if applicable.",
            "Approve the four-part conversation map, episode notes, and transcript structure.",
            "Decide whether full transcripts, edited show notes, or both will publish for SEO and accessibility.",
        ],
    ),
]


def set_cell_or_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side="left", color=GOLD, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, 18, 10, NAVY),
        "Heading 2": (13, 14, 7, BLUE),
        "Heading 3": (12, 10, 5, NAVY),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(BODY)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, size=9, color=GOLD, bold=True)
    r.font.letter_spacing = Pt(1.2)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=30, color=NAVY, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=13.5, color=BODY)
    return p


def add_page_section_title(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=22, color=NAVY, bold=True)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([run_properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_review_box(doc, prompts):
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(6)
    set_paragraph_border(label)
    r = label.add_run("ALEX REVIEW — DECISIONS AND INPUTS")
    set_font(r, size=9, color=NAVY, bold=True)
    for prompt in prompts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(prompt)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def clean_text(element):
    # lxml joins text around <br> without spaces; preserve the visual line break
    # so cover/title copy reads naturally in the review workbook.
    for br in element.xpath(".//br"):
        br.tail = " " + (br.tail or "")
    return " ".join(element.text_content().split())


def is_hidden(element):
    if element.get("aria-hidden") == "true":
        return True
    current = element
    while current is not None:
        classes = set((current.get("class") or "").split())
        if "sr-only" in classes or current.get("hidden") is not None:
            return True
        current = current.getparent()
    return False


def add_current_copy_label(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    set_cell_or_paragraph_shading(p, LIGHT)
    r = p.add_run("CURRENT WEBSITE COPY")
    set_font(r, size=9, color=NAVY, bold=True)


def add_copy_item(doc, tag, text, classes):
    if not text:
        return
    if tag == "h1":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        set_paragraph_border(p, color=GOLD, size="12", space="8")
        r = p.add_run(text)
        set_font(r, size=16, color=NAVY, bold=True)
    elif tag == "h2":
        doc.add_paragraph(text, style="Heading 2")
    elif tag == "h3":
        doc.add_paragraph(text, style="Heading 3")
    elif tag == "h4":
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size=11, color=NAVY, bold=True)
    elif tag == "legend":
        p = doc.add_paragraph()
        r = p.add_run(f"Form question: {text}")
        set_font(r, size=11, color=NAVY, bold=True)
    elif tag == "label":
        doc.add_paragraph(text, style="List Bullet")
    elif tag == "li":
        doc.add_paragraph(text, style="List Bullet")
    elif tag in {"a", "button"}:
        p = doc.add_paragraph()
        r = p.add_run(f"CTA: {text}")
        set_font(r, size=10, color=BLUE, bold=True)
    elif "eyebrow" in classes or "small" in classes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper())
        set_font(r, size=8.5, color=GOLD, bold=True)
    else:
        p = doc.add_paragraph(text)
        if "trust-status" in classes or "story-caption" in classes or "fit-micro" in classes:
            for run in p.runs:
                set_font(run, size=9.5, color=BODY, italic=True)


def extract_copy(path):
    tree = html.fromstring(path.read_text(encoding="utf-8"))
    main = tree.xpath("//main")[0] if tree.xpath("//main") else tree.xpath("//body")[0]
    xpath = (
        ".//*[self::h1 or self::h2 or self::h3 or self::h4 or self::p "
        "or self::legend or self::label or self::li or self::button or self::a "
        "or self::small or self::summary or self::figcaption or self::span or self::strong or self::div]"
    )
    output = []
    for element in main.xpath(xpath):
        if is_hidden(element):
            continue
        tag = element.tag.lower()
        classes = set((element.get("class") or "").split())
        if tag == "div" and not classes.intersection({"serve-card-tag"}):
            continue
        if tag in {"a", "button"}:
            if element.xpath(".//*[self::h1 or self::h2 or self::h3 or self::p]"):
                continue
            if classes.intersection({"rail-button", "hero-pause", "menu-toggle", "menu"}):
                continue
        if tag == "li" and element.xpath(".//*[self::h1 or self::h2 or self::h3 or self::p]"):
            continue
        if tag in {"small", "summary", "figcaption", "span", "strong"}:
            # Parent text-bearing elements already include their inline descendants.
            # Keep standalone labels/microcopy while avoiding duplicate fragments.
            parent = element.getparent()
            duplicate_ancestor = False
            while parent is not None and parent is not main:
                parent_tag = parent.tag.lower() if isinstance(parent.tag, str) else ""
                if parent_tag in {"h1", "h2", "h3", "h4", "p", "legend", "label", "button"}:
                    duplicate_ancestor = True
                    break
                if parent_tag == "li" and not parent.xpath(".//*[self::h1 or self::h2 or self::h3 or self::p]"):
                    duplicate_ancestor = True
                    break
                if parent_tag == "a" and not parent.xpath(".//*[self::h1 or self::h2 or self::h3 or self::p]"):
                    duplicate_ancestor = True
                    break
                parent = parent.getparent()
            if duplicate_ancestor:
                continue
            if element.xpath("./*[self::h1 or self::h2 or self::h3 or self::h4 or self::p]"):
                continue
            if element.xpath("./*[self::small or self::summary or self::span or self::strong]") and "text-link" not in classes:
                continue
        text = clean_text(element)
        if not text:
            continue
        if tag == "span" and text.isdigit() and "chapter-number" not in classes:
            continue
        if output and output[-1][1] == text:
            continue
        output.append((tag, text, classes))
    return output


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("ANCHOR WEALTH PLANNING  |  COPY REVIEW")
    set_font(r, size=8.5, color=BODY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Review draft  |  July 17, 2026  |  ")
    set_font(r, size=8.5, color=BODY)
    add_page_field(footer)

    add_kicker(doc, "Client copy workbook")
    add_title(doc, "Current Website Copy & Alex Review Guide")
    add_subtitle(doc, "A page-by-page record of the current Anchor Wealth Planning review site, with the decisions, facts, and approvals needed from Alex before launch.")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Prepared for: Alex Miller and the Anchor review team\nCurrent public review: https://caebaby.github.io/anchor-review/")
    set_font(r, size=10.5, color=BODY)

    doc.add_paragraph("How to use this document", style="Heading 1")
    for item in [
        "Read the Alex Review bullets first. Those are the decisions or facts most likely to change the site.",
        "The Current Website Copy sections reproduce the meaningful on-page copy from the present review build. Navigation and repeated footer links are summarized once.",
        "Correct wording directly, add comments, or answer in plain language. The final site copy can be rewritten from Alex’s answers.",
        "Anything involving registration, fiduciary status, credentials, prior employment, tax language, awards, or testimonials requires factual and compliance approval before launch.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Launch-critical facts", style="Heading 1")
    add_review_box(
        doc,
        [
            "Exact legal entity name and registration status, including state or SEC registration as applicable.",
            "Whether “fiduciary,” “RIA,” “fee-only,” or “fee-based” may be used, and the approved disclosure language.",
            "Exact business phone, email, office location, scheduling link, and compliance contact.",
            "Alex’s exact title, verified years in the industry, approved prior-firm history, credentials, licenses, and accolades.",
            "Service scope and fee language, including what Anchor provides directly and what is coordinated through outside professionals.",
            "Fit Check audience and asset bands, response workflow, response-time promise, and whether the assessment can make a “fit” recommendation.",
            "Approved photography, Alex story video, team headshots, article authorship, and podcast media/platform links.",
        ],
    )

    doc.add_paragraph("Global site copy", style="Heading 1")
    add_review_box(
        doc,
        [
            "Approve the brand line “Charting the course towards your financial legacy” or replace it.",
            "Confirm the preferred firm description: Houston-based fiduciary wealth management for oil and gas executives, business owners, and high-net-worth families.",
            "Provide the final legal footer and privacy/contact requirements.",
        ],
    )
    add_current_copy_label(doc)
    for text in [
        "Primary navigation: About · Team · Who We Serve · Resources · Process · Take the Fit Check",
        "Brand line: “Charting the course towards your financial legacy.”",
        "Firm description: Houston-based fiduciary wealth management for oil & gas executives, business owners, and high-net-worth families.",
        "Current legal placeholder: [VERIFY: RIA entity name] is a registered investment advisor. Information presented is for educational purposes only and does not intend to make an offer or solicitation for the sale or purchase of any specific securities, investments, or investment strategies. Investments involve risk and unless otherwise stated, are not guaranteed. Past performance is not a guarantee of future results.",
        "Repeated Fit Check footer headline: Four questions. A better first conversation.",
        "Repeated Fit Check footer explainer: See whether your situation, timing, and coordination needs match the way Anchor works.",
        "Repeated Fit Check footer steps: 01 Your situation · 02 Current complexity · 03 Coordination gap · 04 Best next step.",
        "Repeated Fit Check footer CTA: Start the Fit Check.",
    ]:
        doc.add_paragraph(text)

    for title, filename, url, prompts in PAGES:
        divider = doc.add_paragraph()
        divider.paragraph_format.space_before = Pt(18)
        divider.paragraph_format.space_after = Pt(10)
        set_paragraph_border(divider, side="top", color=LINE, size="6", space="5")
        add_kicker(doc, "Page copy")
        add_page_section_title(doc, title)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        add_hyperlink(p, "Open this page in the public review", url)
        add_review_box(doc, prompts)
        add_current_copy_label(doc)

        if filename == "home-v6.html":
            doc.add_paragraph("Rotating hero copy", style="Heading 2")
            hero_variants = [
                ("Oil & Gas", "When your equity vests, who sees the whole picture?", "RSUs, deferred comp, blackout periods, commodity cycles. Anchor coordinates your investments, taxes, and estate plan with your CPA and attorney — so every vest, bonus, and decision happens in context."),
                ("Business Owners", "Your business is your largest asset. What’s the plan for it?", "Entity structure, exit timing, liquidity events, and the line between business value and personal wealth. Anchor starts the planning conversation years before the LOI arrives."),
                ("Families", "Your wealth grew. Did your plan keep up?", "Trust funding, beneficiary reviews, multi-generational coordination. Anchor keeps your CPA, attorney, and advisor reading from the same page — so your estate matches your actual net worth."),
            ]
            for label, headline, subhead in hero_variants:
                p = doc.add_paragraph()
                r = p.add_run(f"{label}: {headline}")
                set_font(r, size=11, color=NAVY, bold=True)
                doc.add_paragraph(subhead)
            doc.add_paragraph("CTA: Schedule a 30-Minute Conversation")
            doc.add_paragraph("CTA: Take the 4-Question Fit Check")
            doc.add_paragraph("Proof labels: [Credential 01] · [Credential 02] · [XX]+ yrs · Houston-based")

        for tag, text, classes in extract_copy(BUILD / filename):
            add_copy_item(doc, tag, text, classes)

    doc.core_properties.title = "Anchor Wealth Planning — Current Website Copy & Alex Review Guide"
    doc.core_properties.subject = "Current review-site copy and client approval questions"
    doc.core_properties.author = "Website Factory"
    doc.core_properties.keywords = "Anchor Wealth Planning, website copy, client review"
    doc.save(OUTPUT)
    return OUTPUT


def markdown_copy_item(tag, text, classes):
    if tag == "h1":
        return f"**H1:** {text}"
    if tag == "h2":
        return f"#### {text}"
    if tag == "h3":
        return f"##### {text}"
    if tag == "h4":
        return f"**SUBHEAD:** {text}"
    if tag == "legend":
        return f"**FORM QUESTION:** {text}"
    if tag == "label":
        return f"- **OPTION:** {text}"
    if tag == "li":
        return f"- {text}"
    if tag in {"a", "button"}:
        return f"- **CTA:** {text}"
    if "eyebrow" in classes or "small" in classes:
        return f"**SECTION LABEL:** {text}"
    return text


def build_ai_packet():
    lines = [
        "# ANCHOR WEALTH PLANNING — AI COPY PACKET",
        "",
        "> Upload this file to an AI and say: **Start the review. Ask me one decision at a time.**",
        "",
        "## Instructions for the AI",
        "",
        "You are helping Alex Miller approve and revise the copy for the Anchor Wealth Planning review website.",
        "",
        "Follow this workflow:",
        "",
        "1. Read the entire packet before responding.",
        "2. Start with **Launch blockers**. Ask Alex one short question at a time.",
        "3. Do not rewrite a section until Alex has answered the related decision questions.",
        "4. After each answer, show only: **Decision captured**, **Affected pages**, and **Recommended final copy**.",
        "5. Preserve approved copy that Alex did not ask to change.",
        "6. Never invent facts, credentials, years, awards, testimonials, registration status, fees, client counts, prior employment, or performance claims.",
        "7. Keep unresolved facts marked `[VERIFY]` and list the exact missing input.",
        "8. Treat tax, legal, investment, fiduciary, RIA, and registration language as requiring compliance approval.",
        "9. Keep the voice calm, precise, direct, Houston-aware, and sophisticated. Avoid hype, generic luxury language, and unsupported superiority claims.",
        "10. When all questions are answered, return a final page-by-page copy deck using this format: `PAGE → SECTION → FINAL COPY → VERIFICATION NOTES`.",
        "",
        "## Launch blockers",
        "",
        "Ask these first, one at a time:",
        "",
        "1. What is the exact legal entity name and registration status?",
        "2. Can the site use the words “fiduciary,” “RIA,” “fee-only,” or “fee-based”?",
        "3. What disclosure language must appear in the footer?",
        "4. What are Alex’s exact title, verified years in the industry, approved prior-firm history, credentials, licenses, and accolades?",
        "5. What are the business phone, email, office location, scheduling link, and compliance contact?",
        "6. Which services does Anchor provide directly, and which are coordinated through outside professionals?",
        "7. What are the approved Fit Check asset bands and follow-up workflow?",
        "8. Is the primary conversion action the Fit Check, a 30-minute call, or the Fit Check followed by scheduling?",
        "9. Who belongs on the Team page, and what approved biography/headshot should be used for each person?",
        "10. Which article, podcast, and video topics can Alex credibly publish first?",
        "",
        "## Global and repeated copy",
        "",
        "### Decisions Alex should make",
        "",
        "- Approve or replace the brand line.",
        "- Confirm the firm description and preferred audience labels.",
        "- Supply the final legal footer and privacy/contact requirements.",
        "- Approve the four-step Fit Check description.",
        "",
        "### Current copy",
        "",
        "**Navigation:** About · Team · Who We Serve · Resources · Process · Take the Fit Check",
        "",
        "**Brand line:** “Charting the course towards your financial legacy.”",
        "",
        "**Firm description:** Houston-based fiduciary wealth management for oil & gas executives, business owners, and high-net-worth families.",
        "",
        "**Legal placeholder:** `[VERIFY: RIA entity name]` is a registered investment advisor. Information presented is for educational purposes only and does not intend to make an offer or solicitation for the sale or purchase of any specific securities, investments, or investment strategies. Investments involve risk and unless otherwise stated, are not guaranteed. Past performance is not a guarantee of future results.",
        "",
        "**Repeated Fit Check headline:** Four questions. A better first conversation.",
        "",
        "**Repeated Fit Check explainer:** See whether your situation, timing, and coordination needs match the way Anchor works.",
        "",
        "**Repeated Fit Check steps:** 01 Your situation · 02 Current complexity · 03 Coordination gap · 04 Best next step.",
        "",
        "**Repeated Fit Check CTA:** Start the Fit Check.",
        "",
    ]

    for title, filename, url, prompts in PAGES:
        lines.extend([
            "---",
            "",
            f"## Page: {title}",
            "",
            f"**Review URL:** {url}",
            "",
            "### Decisions Alex should make",
            "",
        ])
        lines.extend(f"- {prompt}" for prompt in prompts)
        lines.extend(["", "### Current copy", ""])

        if filename == "home-v6.html":
            lines.extend([
                "#### Rotating hero copy",
                "",
                "**Oil & Gas H1:** When your equity vests, who sees the whole picture?",
                "",
                "RSUs, deferred comp, blackout periods, commodity cycles. Anchor coordinates your investments, taxes, and estate plan with your CPA and attorney — so every vest, bonus, and decision happens in context.",
                "",
                "**Business Owners H1:** Your business is your largest asset. What’s the plan for it?",
                "",
                "Entity structure, exit timing, liquidity events, and the line between business value and personal wealth. Anchor starts the planning conversation years before the LOI arrives.",
                "",
                "**Families H1:** Your wealth grew. Did your plan keep up?",
                "",
                "Trust funding, beneficiary reviews, multi-generational coordination. Anchor keeps your CPA, attorney, and advisor reading from the same page — so your estate matches your actual net worth.",
                "",
                "- **CTA:** Schedule a 30-Minute Conversation",
                "- **CTA:** Take the 4-Question Fit Check",
                "- **PROOF:** `[Credential 01]` · `[Credential 02]` · `[XX]+ yrs` · Houston-based",
                "",
            ])

        for tag, text, classes in extract_copy(BUILD / filename):
            rendered = markdown_copy_item(tag, text, classes)
            if rendered:
                lines.extend([rendered, ""])

    lines.extend([
        "---",
        "",
        "## Required final AI output",
        "",
        "When Alex has answered every applicable question, return:",
        "",
        "1. **Launch blockers still open**",
        "2. **Approved facts and claims**",
        "3. **Final page-by-page copy**",
        "4. **Compliance review list**",
        "5. **Asset and media request list**",
        "6. **Change log: current copy → final copy**",
        "",
    ])
    AI_OUTPUT.write_text("\n".join(lines), encoding="utf-8")
    return AI_OUTPUT


if __name__ == "__main__":
    print(build_document())
    print(build_ai_packet())
